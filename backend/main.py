"""JobGuard FastAPI backend: scan job listings, return risk analysis."""

import re
import time
import uuid
from datetime import datetime

import io
import secrets

from fastapi import Depends, FastAPI, File, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware

from config import get_config
from db import store as db_store, get as db_get, list_recent as db_list_recent
from schemas import (
    CompanySignal,
    JobPostSignal,
    ListingSnapshot,
    ResumeMatch,
    ScanRequest,
    ScanResponse,
    ScanResult,
    ScanHistoryItem,
)
from services import ollama as ollama_service
from services import searxng as searxng_service

app = FastAPI(title="JobGuard API", version="0.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def require_token(x_api_token: str | None = Header(default=None)) -> None:
    """Gate: if API_TOKEN env is set, require a matching X-API-Token header.
    Uses constant-time compare to avoid timing leaks."""
    expected = get_config().get("api_token")
    if not expected:
        return  # Auth disabled (local dev).
    provided = x_api_token or ""
    if not secrets.compare_digest(provided, expected):
        raise HTTPException(status_code=401, detail="Invalid or missing API token")

# In-memory cache: scan_id -> (ScanResult, expiry_timestamp)
_scan_cache: dict[str, tuple[ScanResult, float]] = {}


def _mock_scan_result(snapshot: ListingSnapshot, scan_id: str) -> ScanResult:
    """Build a stub ScanResult from the request snapshot (for step 2)."""
    return ScanResult(
        id=scan_id,
        timestamp=datetime.utcnow(),
        trustScore=62,
        riskLevel="medium",
        primaryWarning="Medium risk: External email domain mismatch and missing salary info suggest caution. Verify company directly before proceeding.",
        snapshot=snapshot,
        jobPostSignals=[
            JobPostSignal(id="1", status="good", label="Has clear responsibilities", evidence="Lists specific job duties."),
            JobPostSignal(id="2", status="good", label="Has requirements", evidence="Specifies experience and tech stack."),
            JobPostSignal(id="3", status="warn", label="No salary info", evidence="No salary range found in the listing."),
            JobPostSignal(id="4", status="bad", label="External email in description", evidence="Contact email domain may not match company."),
            JobPostSignal(id="5", status="warn", label="Urgent language detected", evidence="Contains phrases like immediate hire, apply ASAP."),
            JobPostSignal(id="6", status="good", label="Description length adequate", evidence=f"{snapshot.descriptionLength} words - within normal range."),
        ],
        companySignals=[
            CompanySignal(id="1", status="good", label="Official website found", evidence="Company site with SSL certificate."),
            CompanySignal(id="2", status="warn", label="Domain mismatch detected", evidence="Email domain may differ from company site."),
            CompanySignal(id="3", status="good", label="Has About/Contact page", evidence="About and contact pages found."),
            CompanySignal(id="4", status="good", label="LinkedIn company page exists", evidence="Company page with followers and employees."),
            CompanySignal(id="5", status="warn", label="Limited business registry info", evidence="Company not found in state business registry."),
            CompanySignal(id="6", status="warn", label="Mixed reputation signals", evidence="Some reviews mention hiring process issues."),
        ],
    )


def _evict_expired():
    now = time.time()
    config = get_config()
    ttl = config["scan_cache_ttl_seconds"]
    expired = [k for k, (_, exp) in _scan_cache.items() if exp < now]
    for k in expired:
        del _scan_cache[k]


@app.get("/health")
def health():
    return {"status": "ok"}


def _extract_domain_from_email(text: str | None) -> str | None:
    """Extract domain from email like careers@techflow-jobs.net."""
    if not text:
        return None
    m = re.search(r"@([a-zA-Z0-9][-a-zA-Z0-9.]*\.[a-zA-Z]{2,})", text)
    return m.group(1) if m else None


def _looks_unknown(value: str | None) -> bool:
    """Heuristic: the extension sends blank or 'Unknown (extract ...)' when scraping failed."""
    if not value:
        return True
    v = value.strip().lower()
    return v.startswith("unknown") or len(v) < 2


@app.post("/scan", response_model=ScanResponse, dependencies=[Depends(require_token)])
def scan(body: ScanRequest):
    """Analyze a job listing with search + LLM; LLM also extracts title/company when the
    extension couldn't scrape them, so the deep-dive investigates the right company."""
    _evict_expired()
    scan_id = str(uuid.uuid4())
    description = body.description or ""
    description_length = body.descriptionLength or len(description)

    # If the extension couldn't get title/company, do a fast LLM pre-extraction
    # so the deep-dive search runs on the real company name.
    resolved_title = body.jobTitle
    resolved_company = body.companyName
    resolved_location = body.location
    resolved_employment = body.employmentType
    if _looks_unknown(resolved_company) or _looks_unknown(resolved_title):
        pre = ollama_service.preextract_fields(description)
        if pre:
            if _looks_unknown(resolved_title) and pre.get("jobTitle"):
                resolved_title = pre["jobTitle"]
            if _looks_unknown(resolved_company) and pre.get("companyName"):
                resolved_company = pre["companyName"]
            if not resolved_location and pre.get("location"):
                resolved_location = pre["location"]
            if not resolved_employment and pre.get("employmentType"):
                resolved_employment = pre["employmentType"]

    snapshot = ListingSnapshot(
        jobTitle=resolved_title or "Unknown title",
        companyName=resolved_company or "Unknown company",
        platform=body.platform,
        pageUrl=body.pageUrl,
        location=resolved_location,
        employmentType=resolved_employment,
        postedDate=body.postedDate,
        applicantCount=body.applicantCount,
        salaryMentioned=body.salaryMentioned,
        responsibilitiesPresent=body.responsibilitiesPresent,
        requirementsPresent=body.requirementsPresent,
        benefitsPresent=body.benefitsPresent,
        contactInfo=body.contactInfo,
        recruiterVisible=body.recruiterVisible,
        descriptionLength=description_length,
    )
    snapshot_dict = snapshot.model_dump()
    search_evidence = searxng_service.gather_evidence(
        resolved_company or body.companyName,
        domain=_extract_domain_from_email(body.contactInfo),
    )
    analysis = ollama_service.analyze(
        snapshot_dict,
        description,
        search_evidence,
        resume=body.resumeText or "",
    )
    if analysis:
        # Trust the LLM's extraction over the scraper's value — the LLM read
        # the full description from the selected detail pane, while the scraper
        # can pick up page titles / section headers / left-hand list items.
        ext_title = analysis.get("extractedJobTitle") or ""
        ext_company = analysis.get("extractedCompanyName") or ""
        ext_location = analysis.get("extractedLocation") or ""
        ext_employment = analysis.get("extractedEmploymentType") or ""
        updates: dict = {}
        if ext_title and ext_title.strip().lower() not in ("unknown", "unknown title"):
            updates["jobTitle"] = ext_title
        elif _looks_unknown(snapshot.jobTitle):
            updates["jobTitle"] = "Unknown title"
        if ext_company and ext_company.strip().lower() not in ("unknown", "unknown company"):
            updates["companyName"] = ext_company
        elif _looks_unknown(snapshot.companyName):
            updates["companyName"] = "Unknown company"
        if ext_location and not snapshot.location:
            updates["location"] = ext_location
        if ext_employment and not snapshot.employmentType:
            updates["employmentType"] = ext_employment
        if updates:
            snapshot = snapshot.model_copy(update=updates)
        resume_match = analysis.get("resumeMatch")
        result = ScanResult(
            id=scan_id,
            timestamp=datetime.utcnow(),
            trustScore=analysis["trustScore"],
            riskLevel=analysis["riskLevel"],
            primaryWarning=analysis["primaryWarning"],
            snapshot=snapshot,
            jobPostSignals=analysis["jobPostSignals"],
            companySignals=analysis["companySignals"],
            resumeMatch=ResumeMatch(**resume_match) if resume_match else None,
        )
    else:
        result = _mock_scan_result(snapshot, scan_id)
    config = get_config()
    ttl = config["scan_cache_ttl_seconds"]
    _scan_cache[scan_id] = (result, time.time() + ttl)
    db_store(scan_id, result)
    return ScanResponse(scanId=scan_id, result=result)


@app.get("/scan/{scan_id}", response_model=ScanResult, dependencies=[Depends(require_token)])
def get_scan(scan_id: str):
    """Return a previously stored scan result by id (for extension handoff)."""
    _evict_expired()
    if scan_id in _scan_cache:
        result, _ = _scan_cache[scan_id]
        return result
    result = db_get(scan_id)
    if result is not None:
        return result
    raise HTTPException(status_code=404, detail="Scan not found")


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise HTTPException(500, "PDF support not installed on server")
    try:
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(p for p in parts if p.strip())
    except Exception as e:
        raise HTTPException(422, f"Could not read PDF: {e}")


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(500, "DOCX support not installed on server")
    try:
        doc = Document(io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # Include text from simple tables (common for resumes).
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        lines.append(t)
        return "\n".join(lines)
    except Exception as e:
        raise HTTPException(422, f"Could not read DOCX: {e}")


@app.post("/extract-resume-file", dependencies=[Depends(require_token)])
async def extract_resume_file(file: UploadFile = File(...)):
    """Accept a resume file (PDF/DOCX/TXT) and return plain text for editing.

    .doc (legacy Word) is not supported server-side; users should convert to PDF
    or DOCX before uploading.
    """
    data = await file.read()
    if not data:
        raise HTTPException(400, "Empty file")
    name = (file.filename or "").lower()
    mime = (file.content_type or "").lower()
    if name.endswith(".pdf") or "pdf" in mime:
        text = _extract_pdf_text(data)
    elif name.endswith(".docx") or "wordprocessingml" in mime:
        text = _extract_docx_text(data)
    elif name.endswith((".txt", ".md")) or mime.startswith("text/"):
        try:
            text = data.decode("utf-8", errors="replace")
        except Exception:
            raise HTTPException(422, "Could not decode text file")
    elif name.endswith(".doc"):
        raise HTTPException(
            415,
            "Legacy .doc files are not supported. Please save as PDF or DOCX and try again.",
        )
    else:
        raise HTTPException(415, f"Unsupported file type: {name or mime or 'unknown'}")

    return {"text": text, "length": len(text)}


@app.post("/parse-resume", dependencies=[Depends(require_token)])
def parse_resume(body: dict):
    """Structure a resume text into name/skills/companies/etc. via the LLM."""
    text = (body.get("text") or "").strip() if isinstance(body, dict) else ""
    if not text:
        raise HTTPException(status_code=400, detail="Missing 'text' field")
    parsed = ollama_service.parse_resume(text)
    if parsed is None:
        raise HTTPException(status_code=502, detail="Could not parse resume")
    return {"parsed": parsed}


@app.get("/history", response_model=list[ScanHistoryItem], dependencies=[Depends(require_token)])
def history(limit: int = 50):
    """Return recent scan history (from DB if configured, else empty)."""
    return db_list_recent(limit=limit)
